"""Generates the Habit Assistant user-facing documentation (v1.9.4):

  - user_manual_EN.docx / user_manual_TH.docx
  - onboarding_EN.docx  / onboarding_TH.docx
  - onboarding_EN.txt   / onboarding_TH.txt   (Telegram-paste-ready, UTF-8 no BOM)

Every fact in this script (command syntax, Thai aliases, config defaults,
message copy) was cross-checked against the actual v1.9.4 source
(core/commands.py, core/i18n.py, core/discoverability.py, core/habitdef.py,
core/routines.py, core/backfill.py, core/quicklog.py, core/cadence.py,
core/pause.py, core/grace.py, core/wrapped.py, core/dashboard.py,
core/heatmap.py, core/records.py, core/trends.py, core/history_view.py,
core/reminders.py, core/checkins.py, core/nudge.py, core/streaks.py,
core/access.py, core/audit_view.py, config.toml) at git tag v1.9.4 -- not
paraphrased from the SPEC-*.md planning docs, which occasionally drift from
what actually shipped.

Re-run after any future release:
    <path-to-venv>\\Scripts\\python.exe docs\\generate_docs.py
"""

from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

APP_VERSION = "1.9.4"
OUT_DIR = pathlib.Path(__file__).resolve().parent

# ---------------------------------------------------------------------------
# Palette / fonts (per Patty's DOCX formatting standard)
# ---------------------------------------------------------------------------
BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK_GRAY = RGBColor(0x1F, 0x29, 0x37)
MED_GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

TABLE_HEADER_BG = "1A56DB"
TABLE_ALT_BG = "F0F4FF"
TIP_BG = "E8F0FE"
WARN_BG = "FFF3CD"
WARN_BORDER = "F59E0B"
TIP_BORDER = "1A56DB"
GRAY_BORDER = "BFBFBF"

FONT_EN = "Calibri"
FONT_TH = "TH Sarabun New"


def FONT(lang: str) -> str:
    return FONT_TH if lang == "th" else FONT_EN


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------


def set_run_font(run, lang: str, size: int = 11, bold: bool = False, italic: bool = False, color=None) -> None:
    name = FONT(lang)
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def shade_cell(cell, hex_color: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def set_cell_borders(cell, sides: dict) -> None:
    """sides: {'top': (sz_eighths, color) | None, 'bottom': ..., 'left': ..., 'right': ...}"""
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, spec in sides.items():
        el = OxmlElement(f"w:{side}")
        if spec is None:
            el.set(qn("w:val"), "nil")
        else:
            sz, color = spec
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    tcpr.append(borders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcpr.append(mar)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, "en", size=9, color=MED_GRAY)


# ---------------------------------------------------------------------------
# Document scaffolding
# ---------------------------------------------------------------------------


def new_document(lang: str):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = FONT(lang)
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), FONT(lang))
    return doc, section


def add_footer(section, doc_title_en: str, doc_title_th: str, lang: str) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.clear_all()
    usable_width = section.page_width - section.left_margin - section.right_margin
    p.paragraph_format.tab_stops.add_tab_stop(usable_width, alignment=3)  # 3 = WD_TAB_ALIGNMENT.RIGHT
    title = doc_title_th if lang == "th" else doc_title_en
    r1 = p.add_run(title)
    set_run_font(r1, lang, size=9, color=MED_GRAY)
    p.add_run("\t")
    label = "หน้า " if lang == "th" else "Page "
    r2 = p.add_run(label)
    set_run_font(r2, lang, size=9, color=MED_GRAY)
    add_page_number_field(p)

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note = f"สำหรับแอปเวอร์ชัน {APP_VERSION}" if lang == "th" else f"for app version {APP_VERSION}"
    r3 = p2.add_run(note)
    set_run_font(r3, lang, size=8, italic=True, color=MED_GRAY)


# ---------------------------------------------------------------------------
# Content-building helpers
# ---------------------------------------------------------------------------


def title_page(doc, lang: str, title_en, title_th, subtitle_en, subtitle_th, date_str: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(140)
    r = p.add_run(title_th if lang == "th" else title_en)
    set_run_font(r, lang, size=24, bold=True, color=BLUE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(10)
    r2 = p2.add_run(subtitle_th if lang == "th" else subtitle_en)
    set_run_font(r2, lang, size=14, color=DARK_GRAY)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(30)
    ver_txt = f"เวอร์ชัน {APP_VERSION}" if lang == "th" else f"Version {APP_VERSION}"
    r3 = p3.add_run(ver_txt)
    set_run_font(r3, lang, size=12, color=MED_GRAY)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(date_str)
    set_run_font(r4, lang, size=11, color=MED_GRAY)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(6)
    prod = "Habit Assistant"
    r5 = p5.add_run(prod)
    set_run_font(r5, lang, size=11, italic=True, color=MED_GRAY)

    doc.add_page_break()


def h1(doc, text: str, lang: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, lang, size=16, bold=True, color=BLUE)
    return p


def h2(doc, text: str, lang: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, lang, size=13, bold=True, color=DARK_GRAY)
    return p


def h3(doc, text: str, lang: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, lang, size=11, bold=True, italic=True, color=DARK_GRAY)
    return p


def body(doc, text: str, lang: str, size: int = 11, space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run_font(r, lang, size=size, color=DARK_GRAY)
    return p


def bullets(doc, items: list, lang: str, size: int = 11):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.6)
        r = p.add_run("•  ")
        set_run_font(r, lang, size=size, bold=True, color=BLUE)
        r2 = p.add_run(item)
        set_run_font(r2, lang, size=size, color=DARK_GRAY)


def numbered(doc, items: list, lang: str, size: int = 11):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.6)
        r = p.add_run(f"{i}.  ")
        set_run_font(r, lang, size=size, bold=True, color=BLUE)
        r2 = p.add_run(item)
        set_run_font(r2, lang, size=size, color=DARK_GRAY)


def callout(doc, text: str, lang: str, kind: str = "tip"):
    """kind: 'tip' (blue) or 'warning' (amber)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    bg = TIP_BG if kind == "tip" else WARN_BG
    border_color = TIP_BORDER if kind == "tip" else WARN_BORDER
    shade_cell(cell, bg)
    set_cell_borders(
        cell,
        {"left": (24, border_color), "top": None, "bottom": None, "right": None},
    )
    set_cell_margins(cell, top=100, bottom=100, left=200, right=160)
    icon = "⚠️ " if kind == "warning" else "💡 "
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(icon + text)
    set_run_font(r, lang, size=10, italic=True, color=DARK_GRAY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return table


def add_table(doc, headers: list, rows: list, lang: str, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        shade_cell(cell, TABLE_HEADER_BG)
        set_cell_borders(
            cell,
            {
                "top": (4, GRAY_BORDER),
                "bottom": (4, GRAY_BORDER),
                "left": (4, GRAY_BORDER),
                "right": (4, GRAY_BORDER),
            },
        )
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, lang, size=10, bold=True, color=WHITE)

    for row_idx, row in enumerate(rows):
        row_cells = table.add_row().cells
        bg = "FFFFFF" if row_idx % 2 == 0 else TABLE_ALT_BG
        for i, val in enumerate(row):
            cell = row_cells[i]
            shade_cell(cell, bg)
            set_cell_borders(
                cell,
                {
                    "top": (4, GRAY_BORDER),
                    "bottom": (4, GRAY_BORDER),
                    "left": (4, GRAY_BORDER),
                    "right": (4, GRAY_BORDER),
                },
            )
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(val))
            set_run_font(r, lang, size=10, color=DARK_GRAY)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def code_line(doc, text: str, lang: str):
    """A monospace-styled single line for a command example."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = BLUE
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), "Consolas")
    rfonts.set(qn("w:eastAsia"), FONT(lang))
    return p


# ===========================================================================
# USER MANUAL content
# ===========================================================================


def build_user_manual(lang: str, date_str: str) -> Document:
    doc, section = new_document(lang)

    def T(en, th):
        return en if lang == "en" else th

    doc_title = T("Habit Assistant -- User Manual", "Habit Assistant -- คู่มือผู้ใช้")
    title_page(
        doc,
        lang,
        "Habit Assistant",
        "Habit Assistant",
        T("User Manual", "คู่มือผู้ใช้"),
        T("User Manual", "คู่มือผู้ใช้"),
        date_str,
    )
    add_footer(section, "Habit Assistant -- User Manual", "Habit Assistant -- คู่มือผู้ใช้", lang)

    # --- Intro -------------------------------------------------------------
    body(
        doc,
        T(
            "Habit Assistant is a private Telegram bot for your household or small group. "
            "You log daily habits -- water, stretching, a diary entry, or anything you or the "
            "bot's owner adds -- by simply typing to it, in English or Thai. It remembers your "
            "streaks, nudges you gently, and never shares your data with anyone else in the group.",
            "Habit Assistant คือบอทส่วนตัวบน Telegram สำหรับครอบครัวหรือกลุ่มเล็กๆ ของคุณ "
            "ใช้บันทึกกิจกรรมประจำวัน เช่น การดื่มน้ำ ยืดเส้น เขียนไดอารี่ หรือกิจกรรมอื่นที่คุณหรือเจ้าของบอทเพิ่มเข้ามา "
            "เพียงพิมพ์คุยกับบอทได้เลย ทั้งภาษาไทยและอังกฤษ บอทจะจำสตรีคของคุณ คอยกระตุ้นเบาๆ "
            "และข้อมูลของคุณจะเป็นส่วนตัว ไม่ปนกับคนอื่นในกลุ่มเด็ดขาด",
        ),
        lang,
    )
    callout(
        doc,
        T(
            "Every setting in this manual (reminder times, check-in window, streak milestones, "
            "backfill limit, and more) reflects how this bot is actually configured right now -- "
            "not generic defaults.",
            "ทุกค่าที่กล่าวถึงในคู่มือนี้ (เวลาแจ้งเตือน ช่วงเวลาเช็คอิน จำนวนวันของไมล์สโตน ขีดจำกัดการบันทึกย้อนหลัง ฯลฯ) "
            "คือค่าที่ตั้งไว้จริงของบอทตัวนี้ ไม่ใช่ค่าเริ่มต้นทั่วไป",
        ),
        lang,
    )

    # =======================================================================
    h1(doc, T("1. Getting Started", "1. เริ่มต้นใช้งาน"), lang)

    h2(doc, T("Joining the bot", "การเข้าร่วมใช้บอท"), lang)
    body(
        doc,
        T(
            "The bot only replies to people its owner has approved. To join:",
            "บอทจะตอบเฉพาะคนที่เจ้าของบอทอนุมัติแล้วเท่านั้น ขั้นตอนการเข้าร่วมมีดังนี้",
        ),
        lang,
    )
    numbered(
        doc,
        [
            T(
                'Open a chat with the bot on Telegram and send /start (or literally anything -- '
                "the first message you ever send starts this process).",
                'เปิดแชทกับบอทใน Telegram แล้วพิมพ์ /start (หรือพิมพ์อะไรก็ได้ -- '
                "ข้อความแรกที่คุณส่งจะเริ่มขั้นตอนนี้โดยอัตโนมัติ)",
            ),
            T(
                "You'll get a message saying the owner has been notified and you're waiting for "
                "approval.",
                "คุณจะได้รับข้อความแจ้งว่าเจ้าของบอทได้รับแจ้งแล้ว และกำลังรอการอนุมัติ",
            ),
            T(
                "Once the owner approves you, you'll immediately get a welcome message -- you're in, "
                "with your own private logs, streaks, and settings from that point on.",
                "เมื่อเจ้าของบอทอนุมัติแล้ว คุณจะได้รับข้อความต้อนรับทันที -- คุณพร้อมใช้งานแล้ว "
                "พร้อมบันทึก สตรีค และการตั้งค่าส่วนตัวของคุณเองตั้งแต่นาทีนั้น",
            ),
        ],
        lang,
    )
    body(
        doc,
        T(
            'If you message again before being approved, you\'ll just get the same "still '
            'waiting" reply -- no need to keep sending /start.',
            "ถ้าส่งข้อความอีกก่อนได้รับการอนุมัติ ก็จะได้รับข้อความ \"ยังรออยู่\" แบบเดิม "
            "ไม่ต้องพิมพ์ /start ซ้ำ",
        ),
        lang,
    )
    callout(
        doc,
        T(
            "Already approved and just want a reminder of what the bot can do? Send /start any "
            'time for a quick "welcome back" note, or /help for the full list.',
            "ถ้าเป็นสมาชิกอยู่แล้วและอยากดูสรุปสิ่งที่บอททำได้ พิมพ์ /start เมื่อไหร่ก็ได้เพื่อรับข้อความต้อนรับกลับ "
            "หรือพิมพ์ /help เพื่อดูรายการทั้งหมด",
        ),
        lang,
    )

    # =======================================================================
    h1(doc, T("2. Logging Your Day", "2. การบันทึกกิจกรรมประจำวัน"), lang)

    h2(doc, T("Just type it", "พิมพ์ได้เลย"), lang)
    body(
        doc,
        T(
            "There is no special syntax for a normal log -- just tell the bot what you did, in "
            "English or Thai, however feels natural:",
            "ไม่ต้องมีรูปแบบพิเศษสำหรับการบันทึกทั่วไป -- แค่บอกบอทว่าคุณทำอะไรไป "
            "จะเป็นภาษาไทยหรืออังกฤษก็ได้ พิมพ์แบบธรรมชาติได้เลย",
        ),
        lang,
    )
    code_line(doc, "500ml", lang)
    code_line(doc, "น้ำ 500 มล.", lang)
    code_line(doc, "10 min stretch", lang)
    code_line(doc, "ยืดเส้น 20 นาที", lang)
    code_line(doc, T("today was tough but I got through it", "วันนี้เหนื่อยหน่อยแต่ก็ผ่านมาได้"), lang)

    h3(doc, T("Two ways the bot understands you", "สองวิธีที่บอทเข้าใจสิ่งที่คุณพิมพ์"), lang)
    add_table(
        doc,
        [T("Path", "รูปแบบ"), T("When it's used", "ใช้เมื่อไหร่"), T("Notes", "หมายเหตุ")],
        [
            [
                T("Instant (no AI)", "บันทึกทันที (ไม่ใช้ AI)"),
                T('A message that\'s just a number + unit, e.g. "500ml" or "2 แก้ว"', 'ข้อความที่มีแค่ตัวเลข + หน่วย เช่น "500ml" หรือ "2 แก้ว"'),
                T(
                    "Works instantly, even if the AI assistant is briefly offline. No ambiguity "
                    "possible -- it only fires when the whole message is exactly NUMBER + UNIT.",
                    "บันทึกได้ทันที แม้ระบบ AI จะขัดข้องชั่วคราว จะทำงานเฉพาะตอนข้อความทั้งหมด "
                    "เป็นแค่ตัวเลข + หน่วยเท่านั้น จึงไม่มีความกำกวม",
                ),
            ],
            [
                T("AI-assisted", "ใช้ AI ช่วยตีความ"),
                T("Anything else -- free text, diary entries, mixed sentences", "ข้อความอื่นๆ ทั้งหมด -- ข้อความอิสระ ไดอารี่ ประโยคผสม"),
                T(
                    "The bot's local AI reads your message and figures out what you meant. If "
                    "it's genuinely unclear, it asks a quick clarifying question instead of "
                    "guessing.",
                    "AI ของบอทจะอ่านข้อความแล้วตีความให้ ถ้าตีความไม่ออกจริงๆ "
                    "บอทจะถามกลับสั้นๆ แทนการเดา",
                ),
            ],
        ],
        lang,
        col_widths=[3.2, 5.5, 6.8],
    )

    h2(doc, T("The one-tap keyboard: /log", "ปุ่มบันทึกด่วน: /log"), lang)
    body(
        doc,
        T(
            'Send /log (Thai: บันทึก) and the bot pops up a row of tap-to-log buttons, built '
            "from your own habits. Tap once and you're done -- no typing.",
            'พิมพ์ /log (หรือ "บันทึก") บอทจะเปิดแถวปุ่มให้แตะบันทึกได้ทันที '
            "สร้างจากรายการกิจกรรมของคุณเอง แตะครั้งเดียวก็เสร็จ ไม่ต้องพิมพ์เลย",
        ),
        lang,
    )
    bullets(
        doc,
        [
            T(
                "A habit with a daily goal (like water) gets a small ladder of amount buttons "
                "(roughly ¼, ½, and the full goal), or the exact quantities you've set up as "
                'shortcuts (water\'s shipped shortcuts are "glass" = 250ml and "bottle" = 600ml).',
                "กิจกรรมที่มีเป้าหมายรายวัน (เช่น น้ำ) จะมีปุ่มปริมาณให้เลือกประมาณ ¼, ½ และเต็มเป้าหมาย "
                'หรือปุ่มลัดที่ตั้งไว้ (ค่าเริ่มต้นของน้ำคือ "แก้ว" = 250 มล. และ "ขวด" = 600 มล.)',
            ),
            T(
                'A yes/no (boolean) habit gets one "done ✓" button.',
                'กิจกรรมแบบทำ/ไม่ทำ จะมีปุ่มเดียวคือ "เสร็จแล้ว ✓"',
            ),
            T(
                "A habit with no goal and no shortcuts (like stretch, by default) has nothing to "
                'tap a single button for, so it doesn\'t appear in /log -- type it instead.',
                "กิจกรรมที่ไม่มีเป้าหมายและไม่มีปุ่มลัด (เช่น ยืดเส้น ตามค่าเริ่มต้น) "
                "จะไม่มีปุ่มให้กดใน /log ให้พิมพ์บันทึกแทน",
            ),
            T(
                "A free-text habit (like diary) never gets a button -- a tap can't carry a "
                "sentence.",
                "กิจกรรมแบบข้อความอิสระ (เช่น ไดอารี่) จะไม่มีปุ่มเลย เพราะการแตะปุ่มพิมพ์ประโยคไม่ได้",
            ),
        ],
        lang,
    )

    h2(doc, T("Instant reactions", "รีแอคชันทันที"), lang)
    body(
        doc,
        T(
            "When you type a log yourself (not a button tap), the bot reacts to your message "
            "with a small emoji -- 💧 for water, 💪 for anything quantifiable (stretch, or a "
            "numeric/duration custom habit), ✅ for a done/text habit. It's purely decorative; "
            "if it fails for any reason your log and confirmation are unaffected.",
            "เมื่อคุณพิมพ์บันทึกเอง (ไม่ใช่การกดปุ่ม) บอทจะกดรีแอคชันเล็กๆ บนข้อความของคุณ "
            "-- 💧 สำหรับน้ำ, 💪 สำหรับกิจกรรมที่นับปริมาณได้ (ยืดเส้น หรือกิจกรรมที่กำหนดเองแบบตัวเลข/ระยะเวลา), "
            "✅ สำหรับกิจกรรมแบบทำเสร็จ/ข้อความ เป็นแค่ของตกแต่ง ถ้าทำไม่สำเร็จก็ไม่กระทบการบันทึกหรือข้อความยืนยันของคุณ",
        ),
        lang,
    )

    h2(doc, T("Fixing a mistake: Undo and Edit", "แก้ไขข้อผิดพลาด: ยกเลิกและแก้ไข"), lang)
    body(
        doc,
        T(
            "Every confirmation message has an ↩️ Undo button right underneath it -- tap it to "
            "remove that entry instantly. You can also type a command:",
            "ทุกข้อความยืนยันจะมีปุ่ม ↩️ ยกเลิก อยู่ด้านล่าง แตะเพื่อลบรายการนั้นได้ทันที "
            "หรือจะพิมพ์คำสั่งก็ได้:",
        ),
        lang,
    )
    add_table(
        doc,
        [T("Action", "การทำงาน"), T("English", "อังกฤษ"), T("Thai", "ไทย")],
        [
            [T("Undo your last entry", "ยกเลิกรายการล่าสุด"), "/undo, “undo”", "ยกเลิก, ลบ"],
            [
                T("Change the value of your last entry", "แก้ไขค่าของรายการล่าสุด"),
                '"make that 300", "change it to 300"',
                "แก้เป็น 300, แก้ไขล่าสุดเป็น 300",
            ],
        ],
        lang,
        col_widths=[5, 5, 5],
    )
    callout(
        doc,
        T(
            "Undo and Edit always act on your own most recent entry -- nobody else in the group "
            "can undo or edit your logs, and you can't touch theirs.",
            "ยกเลิกและแก้ไขจะทำกับรายการล่าสุดของคุณเองเสมอ -- คนอื่นในกลุ่มยกเลิกหรือแก้ไขบันทึกของคุณไม่ได้ "
            "และคุณก็แก้ของคนอื่นไม่ได้เช่นกัน",
        ),
        lang,
    )

    # =======================================================================
    h1(doc, T("3. Logging for an Earlier Day (Backfill)", "3. บันทึกย้อนหลัง"), lang)
    body(
        doc,
        T(
            "Forgot to log yesterday's water? Just add a date phrase to your message -- at the "
            "very start or the very end of it -- and the bot backdates the entry for you, no AI "
            "needed:",
            "ลืมบันทึกน้ำเมื่อวาน? แค่เติมวลีบอกวันที่ต่อท้าย (หรือขึ้นต้น) ข้อความของคุณ "
            "บอทจะบันทึกย้อนหลังให้อัตโนมัติ โดยไม่ต้องใช้ AI เลย:",
        ),
        lang,
    )
    code_line(doc, "500ml yesterday", lang)
    code_line(doc, "เมื่อวาน น้ำ 500", lang)
    code_line(doc, "diary 2 days ago", lang)
    code_line(doc, "ยืดเส้น 20 นาที 3 วันที่แล้ว", lang)
    code_line(doc, "stretched 15 min on Monday", lang)
    code_line(doc, "น้ำ 500 วันจันทร์", lang)

    add_table(
        doc,
        [T("Recognized phrase", "วลีที่บอทเข้าใจ"), T("English", "อังกฤษ"), T("Thai", "ไทย")],
        [
            [T("Yesterday", "เมื่อวาน"), "yesterday", "เมื่อวาน / เมื่อวานนี้"],
            [T("N days ago", "N วันที่แล้ว"), "3 days ago", "3 วันที่แล้ว / 3 วันก่อน"],
            [T("A weekday (most recent past occurrence)", "วันในสัปดาห์ (ที่ผ่านมาล่าสุด)"), "on monday / last friday", "วันจันทร์ / วันศุกร์"],
        ],
        lang,
        col_widths=[5.5, 5, 5],
    )

    callout(
        doc,
        T(
            "The date phrase has to be at the very start or the very end of your message -- "
            '"diary: yesterday was hard" does NOT backfill (the phrase isn\'t at the edge of the '
            "message), but \"diary 2 days ago\" does. This is deliberate: a missed backfill is "
            "recoverable (just try again), but a wrongly-dated log isn't worth the risk.",
            'วลีบอกวันที่ต้องอยู่ตรงต้นหรือท้ายข้อความเท่านั้น -- "ไดอารี่: เมื่อวานเหนื่อยมาก" '
            'จะไม่ถูกบันทึกย้อนหลัง (เพราะวลีไม่ได้อยู่ริมข้อความ) แต่ "ไดอารี่ 2 วันที่แล้ว" จะบันทึกย้อนหลังได้ '
            "เป็นความตั้งใจของระบบ -- ถ้าบันทึกย้อนหลังพลาดก็แค่ลองใหม่ได้ แต่ถ้าบันทึกผิดวันไปแล้วแก้คืนยากกว่า",
        ),
        lang,
    )

    body(
        doc,
        T(
            "You can only backfill up to 14 days back on this bot; anything further (or a future "
            "date) gets a friendly \"too far\" reply, no entry written. A backfilled entry still "
            "shows the normal confirmation, just with a date prefix, e.g. \"📅 Logged for Mon 18 "
            'Aug — ✅ 500 ml logged...\", and Undo works on it exactly like a normal entry.',
            "บนบอทนี้บันทึกย้อนหลังได้สูงสุด 14 วัน ถ้าย้อนไกลกว่านั้น (หรือเป็นวันในอนาคต) "
            'บอทจะตอบว่า "ย้อนหลังไปไกลเกินไป" และไม่บันทึกอะไรให้ รายการที่บันทึกย้อนหลังจะได้ข้อความยืนยันปกติ '
            'แค่มีคำนำหน้าบอกวันที่ เช่น "📅 บันทึกสำหรับ จ. 18 ส.ค. — ✅ บันทึกน้ำ 500 มล. แล้ว..." '
            "และปุ่ม ↩️ ยกเลิก ก็ใช้ได้เหมือนรายการปกติทุกประการ",
        ),
        lang,
    )

    # =======================================================================
    h1(doc, T("4. Custom Habits", "4. นิสัยที่กำหนดเอง"), lang)
    body(
        doc,
        T(
            "Water, stretch, and diary are the built-in habits, but you can add your own from "
            "chat -- entirely private to you, nobody else in the group sees or shares it. Once "
            "created, a custom habit works everywhere a built-in one does: logging, undo, "
            "/target, /remind, streaks, /dashboard, /history, /heatmap, /records, /trends, "
            "check-ins, everything.",
            "น้ำ ยืดเส้น และไดอารี่ เป็นกิจกรรมพื้นฐานที่มากับบอท แต่คุณสามารถเพิ่มกิจกรรมของคุณเองได้จากแชท "
            "เป็นส่วนตัวของคุณคนเดียว คนอื่นในกลุ่มไม่เห็นและไม่ใช้ร่วมกัน เมื่อสร้างแล้วจะใช้งานได้ทุกที่เหมือนกิจกรรมพื้นฐาน "
            "ทั้งการบันทึก ยกเลิก /target /remind สตรีค /dashboard /history /heatmap /records /trends เช็คอิน และอื่นๆ",
        ),
        lang,
    )

    h2(doc, T("/addhabit -- the grammar", "/addhabit -- รูปแบบคำสั่ง"), lang)
    body(
        doc,
        T(
            'The command uses pipe-separated key=value pairs. "id", "type", and "en" are '
            "required; everything else is optional:",
            'คำสั่งนี้ใช้รูปแบบ key=value คั่นด้วย "|" ต้องมี "id", "type" และ "en" ส่วนที่เหลือใส่หรือไม่ใส่ก็ได้:',
        ),
        lang,
    )
    add_table(
        doc,
        [T("Key", "คีย์"), T("Meaning", "ความหมาย"), T("Required?", "จำเป็นไหม")],
        [
            [
                "id",
                T("Lowercase letters/numbers/underscore, max 32 chars, must be unique (can't reuse a built-in or an already-used id, even an archived one)", "ตัวพิมพ์เล็ก ตัวเลข และ _ เท่านั้น ไม่เกิน 32 ตัวอักษร ห้ามซ้ำกับกิจกรรมพื้นฐานหรือ id ที่เคยใช้แล้ว (แม้จะเก็บเข้าคลังไปแล้ว)"),
                T("Yes", "ใช่"),
            ],
            [
                "type",
                T("numeric, duration, text, or boolean", "numeric, duration, text หรือ boolean"),
                T("Yes", "ใช่"),
            ],
            ["en", T("English label", "ชื่อภาษาอังกฤษ"), T("Yes", "ใช่")],
            ["th", T("Thai label (defaults to \"en\" if omitted)", 'ชื่อภาษาไทย (ถ้าไม่ใส่จะใช้ค่าเดียวกับ "en")'), T("No", "ไม่")],
            [
                "unit",
                T('"en/th" -- e.g. "min/นาที". Required for numeric/duration, not allowed for text/boolean', 'รูปแบบ "en/th" เช่น "min/นาที" จำเป็นสำหรับ numeric/duration ห้ามใส่สำหรับ text/boolean'),
                T("Conditional", "ขึ้นอยู่กับประเภท"),
            ],
            [
                "goal",
                T("A positive number -- daily goal (numeric/duration only)", "ตัวเลขบวก -- เป้าหมายรายวัน (เฉพาะ numeric/duration)"),
                T("No", "ไม่"),
            ],
            [
                "alias",
                T('Shortcut units, e.g. "page:1,chapter:10"', 'หน่วยลัด เช่น "page:1,chapter:10"'),
                T("No", "ไม่"),
            ],
            [
                "cadence",
                T('A weekly target instead of daily, e.g. "3w" (1-7)', 'ตั้งเป้าหมายรายสัปดาห์แทนรายวัน เช่น "3w" (1-7)'),
                T("No", "ไม่"),
            ],
        ],
        lang,
        col_widths=[2.3, 9.4, 3],
    )
    body(doc, T("Worked example:", "ตัวอย่าง:"), lang, space_after=2)
    code_line(doc, "/addhabit id=reading|type=duration|en=reading|th=อ่านหนังสือ|unit=min/นาที|goal=30", lang)
    body(
        doc,
        T(
            'This creates a duration habit called "reading" with a 30-minute daily goal -- log it '
            'with "20 min" or "อ่านหนังสือ 20 นาที" from then on.',
            'สร้างกิจกรรมแบบ duration ชื่อ "reading" มีเป้าหมายรายวัน 30 นาที '
            'จากนั้นบันทึกได้ด้วย "20 min" หรือ "อ่านหนังสือ 20 นาที"',
        ),
        lang,
    )
    body(doc, T("A boolean example (no unit, no goal):", "ตัวอย่างแบบ boolean (ไม่มีหน่วย ไม่มีเป้าหมาย):"), lang, space_after=2)
    code_line(doc, "/addhabit id=meditate|type=boolean|en=meditate|th=นั่งสมาธิ", lang)
    body(doc, T("You can create up to 20 active custom habits.", "สร้างกิจกรรมที่กำหนดเองได้สูงสุด 20 รายการที่ยังใช้งานอยู่"), lang)

    h2(doc, T("/delhabit -- smart delete", "/delhabit -- ลบอัจฉริยะ"), lang)
    code_line(doc, "/delhabit reading", lang)
    code_line(doc, "ลบนิสัย reading", lang)
    body(
        doc,
        T(
            "If you've ever logged that habit, it's archived (hidden from active use, id stays "
            "reserved so you can't accidentally reuse it) rather than deleted outright. If you "
            "never logged a single entry, it's permanently removed and the id is freed for reuse.",
            "ถ้าเคยบันทึกกิจกรรมนั้นมาก่อน จะถูกเก็บเข้าคลัง (ซ่อนจากการใช้งาน แต่ id ยังถูกจองไว้ "
            "ป้องกันการสร้างซ้ำโดยไม่ตั้งใจ) แทนที่จะลบทิ้งไปเลย แต่ถ้ายังไม่เคยบันทึกอะไรเลย "
            "จะถูกลบถาวรและ id นั้นนำกลับมาใช้ใหม่ได้",
        ),
        lang,
    )

    # =======================================================================
    h1(doc, T("5. Routines -- bundle several logs into one tap", "5. กิจวัตร -- รวมหลายบันทึกไว้ในคำสั่งเดียว"), lang)
    body(
        doc,
        T(
            'A routine bundles several habit+value entries into one named shortcut, so a whole '
            '"morning stack" logs in one command or one tap.',
            "กิจวัตรคือการรวมหลายรายการ (กิจกรรม + ค่า) ไว้ในชื่อเดียว ทำให้บันทึก \"ชุดกิจกรรมตอนเช้า\" "
            "ทั้งหมดได้ในคำสั่งเดียวหรือแตะปุ่มเดียว",
        ),
        lang,
    )
    add_table(
        doc,
        [T("Action", "การทำงาน"), T("English", "อังกฤษ"), T("Thai", "ไทย")],
        [
            [
                T("Create", "สร้าง"),
                "/routine morning = water 500, stretch 10",
                "กิจวัตร morning = น้ำ 500, ยืดเส้น 10",
            ],
            [T("List all yours (with a run button each)", "ดูรายการทั้งหมด (มีปุ่มรันให้แต่ละอัน)"), "/routine", T("(no Thai bare-word list -- use /routine)", "(ไม่มีคำสั่งภาษาไทยสำหรับดูรายการ -- ใช้ /routine แทน)")],
            [T("Run it", "รัน"), "/routine morning", "กิจวัตร morning"],
            [T("Delete it", "ลบ"), "/routine delete morning", "กิจวัตร morning ลบ"],
        ],
        lang,
        col_widths=[4, 5.5, 5.5],
    )
    callout(
        doc,
        T(
            'Notice the Thai delete order is reversed from English: "กิจวัตร morning ลบ" (name, '
            'then "delete"), not "ลบ กิจวัตร morning". The routine\'s own name always stays plain '
            "lowercase letters/numbers (a-z, 0-9, _) even in a Thai command -- only the habit "
            "items inside it can be Thai.",
            'สังเกตว่าลำดับคำสั่งลบภาษาไทยจะกลับด้าน: "กิจวัตร morning ลบ" (ชื่อก่อน แล้วตามด้วย "ลบ") '
            'ไม่ใช่ "ลบ กิจวัตร morning" ส่วนชื่อกิจวัตรเองต้องเป็นตัวพิมพ์เล็ก ตัวเลข และ _ เท่านั้นเสมอ '
            "แม้จะใช้คำสั่งภาษาไทย มีแต่รายการกิจกรรมข้างในเท่านั้นที่เป็นภาษาไทยได้",
        ),
        lang,
        kind="warning",
    )
    body(
        doc,
        T(
            "Running a routine logs every item at once and gives you one compact summary "
            '("▶️ morning — logged 500 ml water, 10 min stretch (2 of 2)."). A boolean item in a '
            "routine always logs as done; a free-text item (like diary) is always skipped, since a "
            "routine can't carry a sentence for you. You can create up to 20 routines.",
            'การรันกิจวัตรจะบันทึกทุกรายการพร้อมกันและสรุปให้แบบย่อ (เช่น "▶️ morning — บันทึกแล้ว น้ำ 500 มล., '
            'ยืดเส้น 10 นาที (2 จาก 2)") รายการแบบ boolean ในกิจวัตรจะถูกบันทึกว่า "ทำแล้ว" เสมอ '
            "ส่วนรายการแบบข้อความอิสระ (เช่น ไดอารี่) จะถูกข้ามเสมอ เพราะกิจวัตรพิมพ์ประโยคแทนคุณไม่ได้ "
            "สร้างกิจวัตรได้สูงสุด 20 รายการ",
        ),
        lang,
    )

    # =======================================================================
    h1(doc, T("6. Goals & Cadence", "6. เป้าหมายและความถี่รายสัปดาห์"), lang)

    h2(doc, T("Daily goals -- /target", "เป้าหมายรายวัน -- /target"), lang)
    add_table(
        doc,
        [T("Action", "การทำงาน"), T("Command", "คำสั่ง"), T("Notes", "หมายเหตุ")],
        [
            [T("Set a goal", "ตั้งเป้าหมาย"), "/target water 2000", T('or say it naturally: "from now on I want to drink 2.5L a day" / "ต่อไปอยากดื่มน้ำวันละ 2.5 ลิตร"', 'หรือพิมพ์แบบธรรมชาติ: "ต่อไปอยากดื่มน้ำวันละ 2.5 ลิตร"')],
            [T("View one habit's goal", "ดูเป้าหมายของกิจกรรมหนึ่ง"), "/target water", ""],
            [T("View every goal", "ดูเป้าหมายทั้งหมด"), "/target", ""],
            [T("Reset to the bot's default", "รีเซ็ตกลับค่าเริ่มต้น"), "/target water default", ""],
        ],
        lang,
        col_widths=[4.3, 5, 5.7],
    )

    h2(doc, T("Weekly goals -- /cadence", "เป้าหมายรายสัปดาห์ -- /cadence"), lang)
    body(
        doc,
        T(
            'Some habits ("gym 3x/week") don\'t make sense as a daily streak. /cadence turns a '
            "habit's streak into a weekly one -- rest days no longer break it, as long as you hit "
            "your weekly count:",
            'บางกิจกรรม (เช่น "ยิม 3 ครั้ง/สัปดาห์") ไม่เหมาะกับสตรีครายวัน /cadence '
            "จะเปลี่ยนสตรีคของกิจกรรมนั้นให้เป็นแบบรายสัปดาห์แทน -- วันพักจะไม่ทำให้สตรีคขาด "
            "ตราบใดที่ทำครบจำนวนครั้งต่อสัปดาห์ที่ตั้งไว้:",
        ),
        lang,
    )
    code_line(doc, "/cadence gym 3", lang)
    body(
        doc,
        T(
            'This means "gym, 3 times a week" -- your /habits and /dashboard lines for it switch '
            'to "3×/week · this week 2 of 3", and the streak that used to count days now counts '
            "weeks you hit 3. Clear it with:",
            'หมายถึง "ยิม 3 ครั้งต่อสัปดาห์" -- บรรทัดของกิจกรรมนี้ใน /habits และ /dashboard จะเปลี่ยนเป็น '
            '"3 ครั้ง/สัปดาห์ · สัปดาห์นี้ 2 จาก 3" และสตรีคที่เคยนับเป็นวัน จะเปลี่ยนมานับเป็นสัปดาห์ที่ทำครบ 3 ครั้งแทน '
            "ปิดการใช้งานได้ด้วย:",
        ),
        lang,
    )
    code_line(doc, "/cadence gym off", lang)
    body(
        doc,
        T(
            "N can be 1 to 7. You can also set a cadence right when you create a custom habit "
            '(the /addhabit "cadence=3w" key).',
            "N สามารถตั้งได้ตั้งแต่ 1 ถึง 7 และยังตั้งความถี่รายสัปดาห์ตอนสร้างกิจกรรมใหม่ได้เลยด้วย "
            '(ใช้คีย์ "cadence=3w" ใน /addhabit)',
        ),
        lang,
    )

    # =======================================================================
    h1(doc, T("7. Insights -- how you're really doing", "7. ข้อมูลเชิงลึก -- ดูความคืบหน้าของคุณ"), lang)

    h2(doc, T("Live dashboard -- /dashboard", "แดชบอร์ดสด -- /dashboard"), lang)
    body(
        doc,
        T(
            'Send /dashboard on and the bot pins one "Today" message to your chat that keeps '
            "updating itself, in place, every time you log, undo, or edit -- no need to ask again. "
            "If you ever accidentally unpin or delete it, the bot notices next time it tries to "
            "update and quietly re-creates it. Turn it off with /dashboard off; check the current "
            "state with a bare /dashboard.",
            'พิมพ์ /dashboard on บอทจะปักหมุดข้อความ "วันนี้" ไว้ในแชทของคุณ '
            "ซึ่งจะอัปเดตตัวเองทุกครั้งที่คุณบันทึก ยกเลิก หรือแก้ไข ไม่ต้องเปิดดูใหม่เอง "
            "ถ้าคุณเผลอเลิกปักหมุดหรือลบข้อความนั้นไป บอทจะสังเกตเห็นตอนอัปเดตครั้งถัดไปและสร้างใหม่ให้เงียบๆ "
            "ปิดได้ด้วย /dashboard off ดูสถานะปัจจุบันด้วย /dashboard เฉยๆ",
        ),
        lang,
    )

    h2(doc, T("Consistency picture -- /heatmap", "ปฏิทินความสม่ำเสมอ -- /heatmap"), lang)
    add_table(
        doc,
        [T("Command", "คำสั่ง"), T("Shows", "แสดงอะไร")],
        [
            ["/heatmap", T("A GitHub-style calendar image, one strip per habit, last 12 weeks", "ภาพปฏิทินสไตล์ GitHub หนึ่งแถวต่อกิจกรรม ย้อนหลัง 12 สัปดาห์")],
            ["/heatmap water", T("Just water's own strip", "เฉพาะแถบของกิจกรรมน้ำ")],
            ["/heatmap water 8", T("Water, last 8 weeks", "น้ำ ย้อนหลัง 8 สัปดาห์")],
        ],
        lang,
        col_widths=[4.5, 10.5],
    )
    body(doc, T("Darker cells mean closer to (or past) your goal that day. Thai alias: ปฏิทิน.", "สีเข้มขึ้นแปลว่าใกล้หรือถึงเป้าหมายของวันนั้นมากขึ้น ใช้ภาษาไทยได้ด้วยคำว่า ปฏิทิน"), lang)

    h2(doc, T("Personal bests -- /records", "สถิติส่วนตัว -- /records"), lang)
    body(
        doc,
        T(
            "/records (or /records water) shows your best single day, best week, and longest "
            "streak for each habit. The very first thing you ever log quietly becomes your first "
            'record -- no fanfare -- but the moment you genuinely beat it, you get a "🎉 New '
            'personal best" line right in your log confirmation.',
            "/records (หรือ /records water) แสดงวันที่ดีที่สุด สัปดาห์ที่ดีที่สุด และสตรีคยาวที่สุดของแต่ละกิจกรรม "
            'รายการแรกที่คุณบันทึกจะกลายเป็นสถิติแรกเงียบๆ โดยไม่มีการฉลอง แต่พอคุณทำลายสถิตินั้นได้จริง '
            'จะมีข้อความ "🎉 สถิติใหม่" ขึ้นในข้อความยืนยันการบันทึกทันที',
        ),
        lang,
    )

    h2(doc, T("This week vs. last -- /trends", "เทียบสัปดาห์นี้กับสัปดาห์ที่แล้ว -- /trends"), lang)
    body(
        doc,
        T(
            "/trends (or /trends water) is a quick, no-AI, no-nonsense comparison: this week's "
            "total vs. last week's, the percent change, and a callout if you're on a 2+ week rise "
            "or slide.",
            "/trends (หรือ /trends water) คือการเปรียบเทียบแบบรวดเร็ว ไม่ใช้ AI ไม่ซับซ้อน: "
            "ยอดรวมสัปดาห์นี้เทียบกับสัปดาห์ที่แล้ว เปอร์เซ็นต์การเปลี่ยนแปลง และข้อความเตือนถ้าคุณกำลังขึ้นหรือลงต่อเนื่อง 2 สัปดาห์ขึ้นไป",
        ),
        lang,
    )

    h2(doc, T("Shareable recap -- /wrapped", "การ์ดสรุปที่แชร์ได้ -- /wrapped"), lang)
    body(
        doc,
        T(
            "/wrapped (alias /recap) sends one picture card combining your records, trends, and a "
            'mini consistency heatmap. Bare /wrapped covers the last 4 weeks; /wrapped month covers '
            "the current calendar month so far. Nice to screenshot and share.",
            "/wrapped (หรือ /recap) ส่งการ์ดภาพหนึ่งใบที่รวมสถิติ แนวโน้ม และปฏิทินความสม่ำเสมอย่อไว้ด้วยกัน "
            "พิมพ์ /wrapped เฉยๆ จะสรุป 4 สัปดาห์ล่าสุด ส่วน /wrapped month จะสรุปเดือนปัจจุบันจนถึงวันนี้ "
            "เหมาะสำหรับแคปหน้าจอไปแชร์",
        ),
        lang,
    )

    h2(doc, T("Your own statement -- /history", "รายการย้อนหลังของคุณ -- /history"), lang)
    body(
        doc,
        T(
            "/history (Thai: ย้อนหลัง) lists your own raw entries, newest first -- default 20, up "
            "to 50 at once. Filter by habit and/or count:",
            "/history (หรือ ย้อนหลัง) แสดงรายการดิบของคุณเอง เรียงจากล่าสุดไปเก่าสุด "
            "ค่าเริ่มต้น 20 รายการ สูงสุด 50 รายการ กรองตามกิจกรรมและ/หรือจำนวนได้:",
        ),
        lang,
    )
    code_line(doc, "/history water 10", lang)
    code_line(doc, "ย้อนหลัง น้ำ 10", lang)
    body(
        doc,
        T(
            'A line looks like: "• 08-23 14:03 · 500 ml water · \\"500ml\\"" -- timestamp, what '
            "was recorded, and the exact original message quoted. An undone entry is still shown, "
            'marked "(undone)", so your history is a complete, honest record.',
            '"• 08-23 14:03 · น้ำ 500 มล. · \\"500ml\\"" -- เวลา สิ่งที่บันทึก และข้อความต้นฉบับที่พิมพ์ไว้ '
            'รายการที่ยกเลิกไปแล้วก็ยังแสดงอยู่ แต่มีเครื่องหมาย "(ยกเลิกแล้ว)" กำกับ '
            "ทำให้ประวัติของคุณครบถ้วนและตรงไปตรงมาเสมอ",
        ),
        lang,
    )

    # =======================================================================
    h1(doc, T("8. Staying on Track", "8. คอยติดตามอย่างต่อเนื่อง"), lang)

    h2(doc, T("Reminders -- /remind", "การแจ้งเตือน -- /remind"), lang)
    body(
        doc,
        T(
            "Water and stretch already have reminder times built in "
            "(water: 08:00, 10:30, 13:00, 15:30, 18:00, 20:30; stretch: 11:00, 16:00; diary: "
            "21:30). Make them your own:",
            "น้ำและยืดเส้นมีเวลาแจ้งเตือนตั้งไว้ให้แล้ว "
            "(น้ำ: 08:00, 10:30, 13:00, 15:30, 18:00, 20:30; ยืดเส้น: 11:00, 16:00; ไดอารี่: 21:30) "
            "ปรับให้เป็นของคุณเองได้:",
        ),
        lang,
    )
    add_table(
        doc,
        [T("Action", "การทำงาน"), T("Command", "คำสั่ง")],
        [
            [T("View a habit's reminder times", "ดูเวลาแจ้งเตือนของกิจกรรม"), "/remind water"],
            [T("Set custom times (up to 24)", "ตั้งเวลาของคุณเอง (สูงสุด 24 เวลา)"), "/remind water 08:00 12:00 20:00"],
            [T("Turn off", "ปิด"), "/remind water off"],
            [T("Reset to the bot's default", "รีเซ็ตกลับค่าเริ่มต้น"), "/remind water default"],
        ],
        lang,
        col_widths=[6, 9],
    )
    body(
        doc,
        T(
            "A reminder is automatically skipped once you've already met that day's goal for "
            'that habit -- the bot doesn\'t nag you about something you\'ve finished. Running '
            'late? Reply "snooze" (or "snooze 30" for a custom delay; default 30 minutes) or the '
            'Thai "เลื่อน" / "เลื่อน 30 นาที" to the reminder that just fired.',
            "การแจ้งเตือนจะถูกข้ามอัตโนมัติถ้าคุณทำเป้าหมายของวันนั้นสำเร็จแล้ว "
            'บอทจะไม่กวนใจเรื่องที่คุณทำเสร็จไปแล้ว ถ้ายังไม่พร้อม พิมพ์ "เลื่อน" (หรือ "เลื่อน 30 นาที" '
            'เพื่อกำหนดเวลาเอง ค่าเริ่มต้น 30 นาที) ตอบกลับการแจ้งเตือนที่เพิ่งส่งมาได้เลย',
        ),
        lang,
    )

    h2(doc, T("Hourly check-ins & Do-Not-Disturb", "เช็คอินรายชั่วโมง และช่วงเวลางดรบกวน"), lang)
    body(
        doc,
        T(
            "Check-ins are opt-in and off by default. Turn them on for a gentle nudge once an "
            "hour, on the hour, during your own window (08:00-20:00 by default) -- but only when "
            "you actually have something unmet; if every goal-bearing habit is already done, the "
            "bot stays quiet.",
            "เช็คอินเป็นแบบเปิดเอง ปิดอยู่โดยค่าเริ่มต้น เปิดเพื่อรับการเตือนเบาๆ ทุกชั่วโมง ตรงชั่วโมง "
            "ในช่วงเวลาของคุณเอง (ค่าเริ่มต้น 08:00-20:00) แต่จะเตือนเฉพาะตอนที่ยังมีอะไรค้างอยู่จริงๆ "
            "ถ้าทำครบทุกเป้าหมายแล้ว บอทจะเงียบไว้เฉยๆ",
        ),
        lang,
    )
    add_table(
        doc,
        [T("Action", "การทำงาน"), T("English", "อังกฤษ"), T("Thai", "ไทย")],
        [
            [T("Turn on (default window)", "เปิด (ช่วงเวลาเริ่มต้น)"), "/checkin on", "เช็คอิน on"],
            [T("Set your own window", "ตั้งช่วงเวลาของคุณเอง"), "/checkin 09:00-18:00", "เช็คอิน 09:00-18:00"],
            [T("Turn off", "ปิด"), "/checkin off", "เช็คอิน off"],
            [T("Set quiet hours (no reminders at all)", "ตั้งช่วงเวลางดแจ้งเตือนทั้งหมด"), "/quiet 22:00-07:00", "เงียบ 22:00-07:00"],
            [T("Clear quiet hours", "ล้างช่วงเวลางดแจ้งเตือน"), "/quiet off", "เงียบ off"],
        ],
        lang,
        col_widths=[6, 4.5, 4.5],
    )
    body(
        doc,
        T(
            "/dnd is just another name for /quiet, in case that's the word that comes to mind -- "
            "same effect, same window.",
            "/dnd เป็นอีกชื่อของ /quiet ผลลัพธ์และช่วงเวลาเหมือนกันทุกประการ เผื่อจำคำนี้ได้ง่ายกว่า",
        ),
        lang,
    )

    h2(doc, T('The "almost there" nudge', "การเตือนใจ \"ใกล้ถึงแล้ว\""), lang)
    body(
        doc,
        T(
            "If check-ins are on, you also get one gentle nudge at 20:00 on any day you're 80% "
            "or more of the way to a goal but haven't quite finished it -- folded into a single "
            "message even if several habits qualify, and nothing at all if none do.",
            "ถ้าเปิดเช็คอินไว้ คุณจะได้รับการเตือนใจเบาๆ อีกครั้งตอน 20:00 ในวันที่คุณทำได้ 80% ขึ้นไปของเป้าหมาย "
            "แต่ยังไม่ครบ -- จะรวมเป็นข้อความเดียวแม้จะมีหลายกิจกรรมเข้าเงื่อนไข และจะไม่ส่งอะไรเลยถ้าไม่มีกิจกรรมไหนเข้าเงื่อนไข",
        ),
        lang,
    )

    h2(doc, T("Weekly review & daily summary", "รีวิวรายสัปดาห์ และสรุปประจำวัน"), lang)
    body(
        doc,
        T(
            "Every Sunday at 20:00 you get a weekly review with your stats (and chart images, "
            "when available). A daily summary arrives every night at 21:45 with each habit's "
            "total and current streak.",
            "ทุกวันอาทิตย์เวลา 20:00 คุณจะได้รับรีวิวประจำสัปดาห์พร้อมสถิติ (และภาพกราฟ ถ้ามี) "
            "และทุกคืนเวลา 21:45 จะมีสรุปประจำวันแจ้งยอดรวมและสตรีคปัจจุบันของแต่ละกิจกรรม",
        ),
        lang,
    )

    h2(doc, T("Milestones", "ไมล์สโตน"), lang)
    body(
        doc,
        T(
            "Hit 3, 7, or 30 in a row (days for a normal habit, weeks for a cadence habit) and "
            "you'll get a one-off congratulatory line right in your confirmation -- never a "
            "guilt trip for missing one.",
            "ทำต่อเนื่องครบ 3, 7 หรือ 30 (วันสำหรับกิจกรรมปกติ, สัปดาห์สำหรับกิจกรรมที่ตั้งความถี่รายสัปดาห์) "
            "จะมีข้อความแสดงความยินดีขึ้นในข้อความยืนยันครั้งนั้น -- ไม่มีการตำหนิถ้าพลาดแม้แต่ครั้งเดียว",
        ),
        lang,
    )

    h2(doc, T("Grace day 🛟 -- life happens", "วันผ่อนผัน 🛟 -- บางวันก็มีเหตุ"), lang)
    body(
        doc,
        T(
            "This is fully automatic -- there's no command for it. Once a week at most, if you "
            "genuinely miss a day on a regular (non-cadence) habit that had an active streak "
            "going, the bot quietly protects that streak instead of breaking it, and sends you a "
            'kind note: "🛟 No worries — I used your grace day for water, so your 12-day streak is '
            'safe." No punishment, ever -- and if you log something on your own for that day '
            "later, your real entry always counts instead.",
            "ระบบนี้ทำงานอัตโนมัติทั้งหมด ไม่มีคำสั่งให้พิมพ์ อย่างมากสัปดาห์ละครั้ง ถ้าคุณพลาดไปหนึ่งวันจริงๆ "
            "ในกิจกรรมปกติ (ที่ไม่ได้ตั้งความถี่รายสัปดาห์) ที่กำลังมีสตรีคอยู่ บอทจะปกป้องสตรีคนั้นเงียบๆ แทนที่จะปล่อยให้ขาด "
            'พร้อมข้อความให้กำลังใจ เช่น "🛟 ไม่ต้องห่วงนะ — ฉันใช้สิทธิ์ผ่อนผันของคุณให้น้ำแล้ว '
            'สตรีค 12 วันของคุณยังปลอดภัย" ไม่มีการตำหนิใดๆ และถ้าคุณบันทึกของวันนั้นเองในภายหลัง '
            "รายการจริงของคุณจะถูกใช้แทนเสมอ",
        ),
        lang,
    )

    h2(doc, T("Pause / vacation mode", "โหมดพัก / ลาพัก"), lang)
    body(
        doc,
        T(
            "Going away? Pause a habit (or everything) so reminders, check-ins, and nudges stay "
            "quiet and your streak is held, not broken, while you're gone:",
            "จะไปพัก? พักกิจกรรมใดกิจกรรมหนึ่ง (หรือทุกอย่าง) เพื่อให้การแจ้งเตือน เช็คอิน และการกระตุ้นเงียบไว้ชั่วคราว "
            "และคงสตรีคไว้ ไม่ให้ขาดระหว่างที่คุณไม่อยู่:",
        ),
        lang,
    )
    add_table(
        doc,
        [T("Action", "การทำงาน"), T("English", "อังกฤษ"), T("Thai", "ไทย")],
        [
            [T("Pause one habit for N days", "พักกิจกรรมหนึ่งเป็นจำนวนวัน"), "/pause water 5d", "พัก น้ำ 5d"],
            [T("Pause one habit until a date/weekday", "พักกิจกรรมหนึ่งจนถึงวันที่/วันในสัปดาห์"), "/pause water until 2026-09-01", "พัก น้ำ until 2026-09-01"],
            [T("Pause everything", "พักทุกอย่าง"), "/pause 5d", T('(no bare Thai "pause everything" word yet -- use the English /pause form)', '(ยังไม่มีคำภาษาไทยสำหรับ "พักทุกอย่าง" แบบเปล่าๆ -- ใช้รูปแบบภาษาอังกฤษ /pause แทน)')],
            [T("Check what's currently paused", "ดูว่าตอนนี้พักอะไรอยู่"), "/pause", ""],
            [T("Resume one habit early", "กลับมาก่อนกำหนดเฉพาะกิจกรรมหนึ่ง"), "/resume water", "กลับมา น้ำ"],
            [T("Resume everything", "กลับมาทั้งหมด"), "/resume", ""],
        ],
        lang,
        col_widths=[5.5, 5, 4.5],
    )
    body(
        doc,
        T(
            "A pause can be at most 30 days. You can still log manually while paused -- it's "
            'still recorded and still celebrates -- pausing only mutes the PROACTIVE side '
            "(reminders/check-ins/nudges). A paused habit shows a ⏸ marker with the date it "
            "resumes on your dashboard and /habits.",
            "พักได้สูงสุด 30 วัน ระหว่างพักคุณยังบันทึกเองได้ปกติ -- จะถูกบันทึกและฉลองตามปกติ "
            "การพักจะปิดแค่ฝั่งที่บอทเป็นฝ่ายส่งหาคุณ (การแจ้งเตือน/เช็คอิน/การกระตุ้น) "
            "กิจกรรมที่พักอยู่จะมีเครื่องหมาย ⏸ พร้อมวันที่จะกลับมา แสดงในแดชบอร์ดและ /habits",
        ),
        lang,
    )

    # =======================================================================
    h1(doc, T("9. Preferences", "9. การตั้งค่าส่วนตัว"), lang)
    add_table(
        doc,
        [T("Setting", "การตั้งค่า"), T("Command", "คำสั่ง"), T("Notes", "หมายเหตุ")],
        [
            [
                T("Reply language", "ภาษาที่บอทใช้ตอบ"),
                "/lang en | /lang th | /lang auto",
                T('"auto" (the default) matches whichever language you just wrote in; ภาษา en|th|auto works too', '"auto" (ค่าเริ่มต้น) จะตอบตามภาษาที่คุณเพิ่งพิมพ์มา หรือใช้ "ภาษา en|th|auto" ก็ได้'),
            ],
        ],
        lang,
        col_widths=[3.5, 6.5, 5],
    )
    callout(
        doc,
        T(
            "Since v1.8, reminders, check-ins, and nudges arrive silently by default -- the bot "
            "still sends them, they just don't trigger a notification ping (no more buzzing your "
            "phone every 08:00). Log confirmations and things you asked for directly still notify "
            "normally. If this bot's owner has changed that setting, silent behavior may differ.",
            "ตั้งแต่ v1.8 การแจ้งเตือน เช็คอิน และการกระตุ้นจะส่งแบบไม่มีเสียงแจ้งเตือนโดยค่าเริ่มต้น "
            "บอทยังส่งข้อความตามปกติ เพียงแต่ไม่ทำให้มือถือสั่นหรือมีเสียงทุก 08:00 "
            "ส่วนข้อความยืนยันการบันทึกและสิ่งที่คุณขอเองยังคงแจ้งเตือนตามปกติ "
            "ถ้าเจ้าของบอทเปลี่ยนการตั้งค่านี้ พฤติกรรมอาจแตกต่างไปจากนี้",
        ),
        lang,
    )

    # =======================================================================
    h1(doc, T("10. For the Owner", "10. สำหรับเจ้าของบอท"), lang)
    body(
        doc,
        T(
            "These commands only work for the bot's owner -- they're hidden from everyone else's "
            "command menu entirely.",
            "คำสั่งเหล่านี้ใช้ได้เฉพาะเจ้าของบอทเท่านั้น และจะถูกซ่อนจากเมนูคำสั่งของคนอื่นโดยสิ้นเชิง",
        ),
        lang,
    )

    h2(doc, T("Inviting & managing members", "เชิญและจัดการสมาชิก"), lang)
    body(
        doc,
        T(
            "A new person has to message the bot first (anything, even /start) so their chat ID "
            "gets captured -- you'll then get a notification with the exact command to run:",
            "คนใหม่ต้องส่งข้อความหาบอทก่อน (พิมพ์อะไรก็ได้ แม้แต่ /start) เพื่อให้ระบบจับ chat ID ไว้ "
            "จากนั้นคุณจะได้รับแจ้งเตือนพร้อมคำสั่งที่ต้องพิมพ์เพื่ออนุมัติ:",
        ),
        lang,
    )
    code_line(
        doc,
        T(
            "🔔 Jane (chat 123456789) asked for access. Approve with: /approve 123456789",
            "🔔 Jane (แชท 123456789) ขอสิทธิ์เข้าใช้งาน อนุมัติด้วย: /approve 123456789",
        ),
        lang,
    )
    add_table(
        doc,
        [T("Action", "การทำงาน"), T("Command", "คำสั่ง")],
        [
            [T("Approve someone (same as /invite)", "อนุมัติสมาชิก (เหมือนกับ /invite)"), "/approve 123456789"],
            [T("Block someone", "บล็อกสมาชิก"), "/block 123456789"],
            [T("List everyone + role/status/language", "ดูรายชื่อทั้งหมด + สถานะ/บทบาท/ภาษา"), "/users"],
        ],
        lang,
        col_widths=[7.5, 7.5],
    )

    h2(doc, T("Audit log -- /audit", "ประวัติการเปลี่ยนแปลง -- /audit"), lang)
    body(
        doc,
        T(
            "/audit (default last 20, up to 50: /audit 50) shows every account-level change "
            "anyone made -- who, when, what, and old→new -- for accountability. Thai alias: "
            "ประวัติ. Rows are kept for 365 days by default, then pruned automatically.",
            "/audit (ค่าเริ่มต้นแสดง 20 รายการล่าสุด สูงสุด 50: /audit 50) แสดงทุกการเปลี่ยนแปลงระดับบัญชีที่ใครก็ตามทำไว้ "
            "-- ใคร เมื่อไหร่ อะไร และค่าเก่า→ค่าใหม่ -- เพื่อให้ตรวจสอบย้อนหลังได้ ใช้ภาษาไทยได้ด้วยคำว่า ประวัติ "
            "ระบบจะเก็บข้อมูลไว้ 365 วันโดยค่าเริ่มต้น แล้วลบทิ้งอัตโนมัติหลังจากนั้น",
        ),
        lang,
    )

    h2(doc, T("Release announcements", "ประกาศอัปเดตเวอร์ชันใหม่"), lang)
    body(
        doc,
        T(
            "Whenever the bot is updated, every active member gets a one-time \"what's new\" "
            "message, in their own language, the next time they interact with the bot -- nobody "
            "misses an update, and nobody gets spammed twice for the same version.",
            "ทุกครั้งที่บอทได้รับการอัปเดต สมาชิกที่ใช้งานอยู่ทุกคนจะได้รับข้อความ \"มีอะไรใหม่\" ครั้งเดียว "
            "เป็นภาษาของตัวเอง ในครั้งถัดไปที่ใช้งานบอท ไม่มีใครพลาดอัปเดต และไม่มีใครได้รับข้อความซ้ำสำหรับเวอร์ชันเดียวกัน",
        ),
        lang,
    )

    h2(doc, T("Backups", "การสำรองข้อมูล"), lang)
    body(
        doc,
        T(
            "The bot's data is backed up automatically, with the last 14 backups kept on disk. "
            "Restoring a backup or moving the bot to a new machine is an IT-admin task -- ask "
            "whoever set the bot up.",
            "ข้อมูลของบอทจะถูกสำรองอัตโนมัติ โดยเก็บไฟล์สำรอง 14 ครั้งล่าสุดไว้ในเครื่อง "
            "การกู้คืนข้อมูลหรือย้ายบอทไปเครื่องใหม่เป็นงานของฝ่าย IT -- ให้สอบถามคนที่ติดตั้งบอทให้",
        ),
        lang,
    )

    # =======================================================================
    doc.add_page_break()
    h1(doc, T("Appendix A -- Message Syntax Cheat-Sheet", "ภาคผนวก ก -- สรุปรูปแบบข้อความทั้งหมด"), lang)
    body(
        doc,
        T(
            "The bot understands either language for every one of these, regardless of what "
            "language it replies in. Anywhere you see \"habit\", substitute the habit's id or its "
            "label in either language (e.g. water / น้ำ).",
            "บอทเข้าใจทั้งสองภาษาสำหรับทุกคำสั่งด้านล่าง ไม่ว่าจะตอบเป็นภาษาอะไร "
            'ตรงที่เห็นคำว่า "habit" ให้แทนด้วย id หรือชื่อกิจกรรมภาษาใดก็ได้ (เช่น water หรือ น้ำ)',
        ),
        lang,
    )

    cheat_rows = [
        [T("Log something", "บันทึกกิจกรรม"), "500ml / 10 min stretch", "น้ำ 500 มล. / ยืดเส้น 10 นาที"],
        [T("One-tap keyboard", "ปุ่มบันทึกด่วน"), "/log", "บันทึก"],
        [T("Undo last entry", "ยกเลิกรายการล่าสุด"), "/undo", "ยกเลิก, ลบ"],
        [T("Edit last entry", "แก้ไขรายการล่าสุด"), "make that 300 / change it to 300", "แก้เป็น 300"],
        [T("Backfill (yesterday)", "บันทึกย้อนหลัง (เมื่อวาน)"), "500ml yesterday", "เมื่อวาน น้ำ 500"],
        [T("Backfill (N days ago)", "บันทึกย้อนหลัง (N วันที่แล้ว)"), "3 days ago diary...", "3 วันที่แล้ว ..."],
        [T("Backfill (weekday)", "บันทึกย้อนหลัง (วันในสัปดาห์)"), "on monday", "วันจันทร์"],
        [T("Snooze a reminder", "เลื่อนการแจ้งเตือน"), "snooze / snooze 30", "เลื่อน / เลื่อน 30 นาที"],
        [T("Ask a question", "ถามคำถาม"), "how much water this week?", "อาทิตย์นี้ดื่มน้ำไปเท่าไหร่?"],
        [T("Set a daily goal", "ตั้งเป้าหมายรายวัน"), "/target water 2000", '"ต่อไปอยากดื่มน้ำวันละ 2 ลิตร"'],
        [T("View/clear a goal", "ดู/ล้างเป้าหมาย"), "/target water / /target water default", ""],
        [T("Weekly cadence goal", "เป้าหมายรายสัปดาห์"), "/cadence gym 3", "/cadence gym 3"],
        [T("Add a custom habit", "เพิ่มกิจกรรมของตัวเอง"), "/addhabit id=...|type=...|en=...", "เพิ่มนิสัย id=...|type=...|en=..."],
        [T("Remove a custom habit", "ลบกิจกรรมที่กำหนดเอง"), "/delhabit reading", "ลบนิสัย reading"],
        [T("Create a routine", "สร้างกิจวัตร"), "/routine morning = water 500, stretch 10", "กิจวัตร morning = น้ำ 500, ยืดเส้น 10"],
        [T("Run / delete a routine", "รัน / ลบกิจวัตร"), "/routine morning / /routine delete morning", "กิจวัตร morning / กิจวัตร morning ลบ"],
        [T("Reminder times", "เวลาแจ้งเตือน"), "/remind water 08:00 12:00", "เตือน water 08:00 12:00"],
        [T("Reply language", "ภาษาที่ตอบ"), "/lang en / th / auto", "ภาษา en / th / auto"],
        [T("Quiet hours / DND", "ช่วงเวลางดแจ้งเตือน"), "/quiet 22:00-07:00 / /quiet off", "เงียบ 22:00-07:00 / เงียบ off"],
        [T("Hourly check-ins", "เช็คอินรายชั่วโมง"), "/checkin on / /checkin off", "เช็คอิน on / เช็คอิน off"],
        [T("Live dashboard", "แดชบอร์ดสด"), "/dashboard on / off", "แดชบอร์ด on / off"],
        [T("Consistency picture", "ปฏิทินความสม่ำเสมอ"), "/heatmap [habit] [weeks]", "ปฏิทิน [habit] [weeks]"],
        [T("Personal bests", "สถิติส่วนตัว"), "/records [habit]", "สถิติ [habit]"],
        [T("Week-over-week trend", "แนวโน้มรายสัปดาห์"), "/trends [habit]", "แนวโน้ม [habit]"],
        [T("Shareable recap card", "การ์ดสรุป"), "/wrapped [month] / /recap", "การ์ดสรุป / สรุปเดือน"],
        [T("Your own entry list", "รายการบันทึกของคุณ"), "/history [habit] [N]", "ย้อนหลัง [habit] [N]"],
        [T("Pause / resume", "พัก / กลับมา"), "/pause water 5d / /resume water", "พัก น้ำ 5d / กลับมา น้ำ"],
        [T("Your habits, today's progress", "รายการกิจกรรมและความคืบหน้าวันนี้"), "/habits", "นิสัย"],
        [T("Full capability list", "รายการความสามารถทั้งหมด"), "/help", "ช่วยเหลือ, วิธีใช้"],
    ]
    add_table(
        doc,
        [T("What you want", "สิ่งที่ต้องการทำ"), T("English example", "ตัวอย่างภาษาอังกฤษ"), T("Thai example", "ตัวอย่างภาษาไทย")],
        cheat_rows,
        lang,
        col_widths=[4.5, 5.5, 5],
    )

    return doc


# ===========================================================================
# ONBOARDING ONE-PAGER content
# ===========================================================================


def build_onboarding_docx(lang: str, date_str: str) -> Document:
    doc, section = new_document(lang)

    def T(en, th):
        return en if lang == "en" else th

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(T("👋 You're invited to Habit Assistant!", "👋 คุณได้รับเชิญให้ใช้ Habit Assistant แล้ว!"))
    set_run_font(r, lang, size=22, bold=True, color=BLUE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(T(f"Onboarding · v{APP_VERSION} · {date_str}", f"คู่มือเริ่มต้นใช้งาน · v{APP_VERSION} · {date_str}"))
    set_run_font(r2, lang, size=10, italic=True, color=MED_GRAY)

    body(
        doc,
        T(
            "A private Telegram bot for logging little daily habits -- water, stretching, a "
            "diary entry, or anything else -- and gently keeping each other on track. It "
            "understands plain English or Thai. No app to install.",
            "บอทส่วนตัวบน Telegram สำหรับบันทึกกิจกรรมเล็กๆ ประจำวัน เช่น การดื่มน้ำ ยืดเส้น เขียนไดอารี่ หรืออื่นๆ "
            "และคอยเป็นกำลังใจให้กันและกันเบาๆ เข้าใจได้ทั้งภาษาไทยและอังกฤษแบบธรรมชาติ ไม่ต้องติดตั้งแอปเพิ่ม",
        ),
        lang,
        size=12,
        space_after=10,
    )

    h2(doc, T("⏱ Getting started (20 seconds)", "⏱ เริ่มต้นใช้งาน (20 วินาที)"), lang)
    numbered(
        doc,
        [
            T("Open this chat and send /start", "เปิดแชทนี้แล้วพิมพ์ /start"),
            T(
                "Wait for approval -- you'll get a message the moment you're in, usually quick",
                "รอการอนุมัติ -- พอเข้าได้จะมีข้อความแจ้งทันที ปกติใช้เวลาไม่นาน",
            ),
            T(
                "You're set! Every log, streak, and reminder from here is yours alone",
                "พร้อมแล้ว! การบันทึก สตรีค และการแจ้งเตือนทุกอย่างจากนี้เป็นของคุณคนเดียว",
            ),
        ],
        lang,
        size=12,
    )

    h2(doc, T("✨ Try these 3 things", "✨ ลอง 3 อย่างนี้ดู"), lang)
    bullets(
        doc,
        [
            T('💧 Log something -- just type "500ml" or "10 min stretch"', '💧 บันทึกอะไรสักอย่าง -- พิมพ์ "500ml" หรือ "ยืดเส้น 10 นาที" ได้เลย'),
            T("❓ Send /help to see everything the bot can do", "❓ พิมพ์ /help เพื่อดูทุกอย่างที่บอททำได้"),
            T("📋 Send /habits to see what's tracked and how today's going", "📋 พิมพ์ /habits เพื่อดูรายการกิจกรรมและความคืบหน้าวันนี้"),
        ],
        lang,
        size=12,
    )

    callout(
        doc,
        T(
            "Made a mistake? Every confirmation has its own ↩️ Undo button -- no harm done.",
            "พิมพ์ผิดไปนิด ไม่ต้องกังวล -- ทุกข้อความยืนยันมีปุ่ม ↩️ ยกเลิก ให้กดได้เลย",
        ),
        lang,
    )

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(10)
    r3 = p3.add_run(T("That's it -- see you in the chat! 🎉", "แค่นี้เอง -- แล้วเจอกันในแชทนะ! 🎉"))
    set_run_font(r3, lang, size=13, bold=True, color=BLUE)

    add_footer(section, "Habit Assistant -- Onboarding", "Habit Assistant -- คู่มือเริ่มต้นใช้งาน", lang)
    return doc


def onboarding_txt(lang: str) -> str:
    if lang == "en":
        return (
            "👋 You're invited to Habit Assistant!\n\n"
            "A private Telegram bot for logging little daily habits — water, stretching, a "
            "diary entry, or anything else — and gently keeping each other on track. It "
            "understands plain English or Thai. No app to install.\n\n"
            "⏱ Getting started (20 seconds)\n"
            "1️⃣ Open this chat and send /start\n"
            "2️⃣ Wait for approval — you'll get a message the moment you're in, usually quick\n"
            "3️⃣ You're set! Every log, streak, and reminder from here is yours alone\n\n"
            "✨ Try these 3 things\n"
            '💧 Log something — just type "500ml" or "10 min stretch"\n'
            "❓ Send /help to see everything the bot can do\n"
            "📋 Send /habits to see what's tracked and how today's going\n\n"
            "Made a mistake? Every confirmation has its own ↩️ Undo button — no harm done.\n\n"
            "That's it — see you in the chat! 🎉\n\n"
            f"— Habit Assistant v{APP_VERSION}"
        )
    return (
        "👋 คุณได้รับเชิญให้ใช้ Habit Assistant แล้ว!\n\n"
        "บอทส่วนตัวบน Telegram สำหรับบันทึกกิจกรรมเล็กๆ ประจำวัน เช่น การดื่มน้ำ ยืดเส้น "
        "เขียนไดอารี่ หรืออื่นๆ และคอยเป็นกำลังใจให้กันและกันเบาๆ เข้าใจได้ทั้งภาษาไทยและอังกฤษ"
        "แบบธรรมชาติ ไม่ต้องติดตั้งแอปเพิ่ม\n\n"
        "⏱ เริ่มต้นใช้งาน (20 วินาที)\n"
        "1️⃣ เปิดแชทนี้แล้วพิมพ์ /start\n"
        "2️⃣ รอการอนุมัติ — พอเข้าได้จะมีข้อความแจ้งทันที ปกติใช้เวลาไม่นาน\n"
        "3️⃣ พร้อมแล้ว! การบันทึก สตรีค และการแจ้งเตือนทุกอย่างจากนี้เป็นของคุณคนเดียว\n\n"
        "✨ ลอง 3 อย่างนี้ดู\n"
        '💧 บันทึกอะไรสักอย่าง — พิมพ์ "500ml" หรือ "ยืดเส้น 10 นาที" ได้เลย\n'
        "❓ พิมพ์ /help เพื่อดูทุกอย่างที่บอททำได้\n"
        "📋 พิมพ์ /habits เพื่อดูรายการกิจกรรมและความคืบหน้าวันนี้\n\n"
        "พิมพ์ผิดไปนิด ไม่ต้องกังวล — ทุกข้อความยืนยันมีปุ่ม ↩️ ยกเลิก ให้กดได้เลย\n\n"
        "แล้วเจอกันในแชทนะ! 🎉\n\n"
        f"— Habit Assistant v{APP_VERSION}"
    )


# ===========================================================================
# main
# ===========================================================================


_TH_MONTHS = [
    "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
    "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม",
]


def main() -> None:
    import datetime

    today = datetime.date.today()
    date_en = today.strftime("%B %d, %Y")
    # Python's strftime has no Thai locale available on a typical Windows box
    # without installing one -- build the Thai month name ourselves rather
    # than silently falling back to English (Gregorian year, matching the
    # Gregorian dates used in the manual's own backfill examples).
    date_th = f"{today.day} {_TH_MONTHS[today.month - 1]} {today.year}"

    manual_en = build_user_manual("en", date_en)
    manual_en.save(OUT_DIR / "user_manual_EN.docx")

    manual_th = build_user_manual("th", date_th)
    manual_th.save(OUT_DIR / "user_manual_TH.docx")

    onb_en = build_onboarding_docx("en", date_en)
    onb_en.save(OUT_DIR / "onboarding_EN.docx")

    onb_th = build_onboarding_docx("th", date_th)
    onb_th.save(OUT_DIR / "onboarding_TH.docx")

    (OUT_DIR / "onboarding_EN.txt").write_text(onboarding_txt("en"), encoding="utf-8", newline="\n")
    (OUT_DIR / "onboarding_TH.txt").write_text(onboarding_txt("th"), encoding="utf-8", newline="\n")

    print("Generated:")
    for f in sorted(OUT_DIR.glob("*.docx")) + sorted(OUT_DIR.glob("*.txt")):
        print(f"  {f.name}  ({f.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
